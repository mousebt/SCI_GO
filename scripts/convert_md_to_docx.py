import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_styled_document(md_path, docx_path):
    doc = Document()

    # Page setup (A4 with standard margins)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(4)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []
    in_code_block = False
    code_block_lines = []

    def format_text_with_inline_md(paragraph, text, base_bold=False, base_color=None):
        # Parse inline bold **text**, `code`, etc.
        pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\$[^\$]+\$)'
        tokens = re.split(pattern, text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
                if base_color:
                    run.font.color.rgb = base_color
            elif token.startswith('`') and token.endswith('`'):
                run = paragraph.add_run(token[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(180, 40, 40)
            elif token.startswith('$') and token.endswith('$'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(token)
                run.bold = base_bold
                if base_color:
                    run.font.color.rgb = base_color

    def flush_table(rows):
        if not rows:
            return
        parsed_rows = []
        for r in rows:
            if re.match(r'^\s*\|?\s*[-:]+[-| :]*\s*$', r):
                continue
            cols = [c.strip() for c in r.strip().strip('|').split('|')]
            parsed_rows.append(cols)
        
        if not parsed_rows:
            return

        max_cols = max(len(r) for r in parsed_rows)
        table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, row in enumerate(parsed_rows):
            is_header = (i == 0)
            for j, cell_text in enumerate(row):
                if j < max_cols:
                    cell = table.cell(i, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    
                    if is_header:
                        set_cell_background(cell, "1F4E79")
                        format_text_with_inline_md(p, cell_text, base_bold=True, base_color=RGBColor(255, 255, 255))
                    else:
                        bg_color = "F2F5F8" if i % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg_color)
                        format_text_with_inline_md(p, cell_text)
                    set_cell_margins(cell, 80, 80, 120, 120)

        # Space after table
        sp_p = doc.add_paragraph()
        sp_p.paragraph_format.space_after = Pt(6)

    def flush_code_block(code_lines):
        if not code_lines:
            return
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F4F6F9")
        set_cell_margins(cell, 100, 100, 150, 150)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        for cl in code_lines:
            run = p.add_run(cl + "\n")
            run.font.name = 'Consolas'
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(45, 55, 72)
        sp_p = doc.add_paragraph()
        sp_p.paragraph_format.space_after = Pt(6)

    for line in lines:
        raw_line = line.rstrip('\r\n')
        stripped = raw_line.strip()

        # Handle Code Block
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
                flush_code_block(code_block_lines)
                code_block_lines = []
            else:
                if in_table:
                    in_table = False
                    flush_table(table_rows)
                    table_rows = []
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(raw_line)
            continue

        # Handle Table
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            continue
        else:
            if in_table:
                in_table = False
                flush_table(table_rows)
                table_rows = []

        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 78, 121) # Deep Navy
            run.bold = True
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(46, 117, 182) # Accent Blue
            run.bold = True
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(60, 60, 60)
            run.bold = True
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            format_text_with_inline_md(p, stripped[2:].strip())
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            text_part = re.sub(r'^\d+\.\s', '', stripped)
            format_text_with_inline_md(p, text_part)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif stripped.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D3D3D3"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
        else:
            p = doc.add_paragraph()
            format_text_with_inline_md(p, stripped)

    if in_table:
        flush_table(table_rows)
    if in_code_block:
        flush_code_block(code_block_lines)

    doc.save(docx_path)
    print(f"[Success] Generated Word document: {docx_path}")

if __name__ == '__main__':
    md_file = r'e:\pythonProjects\SCI_GO\docs\research-plan-sap-nested-case-control.md'
    docx_file_en = r'e:\pythonProjects\SCI_GO\docs\SAP_Nested_Case_Control_Research_Plan.docx'
    docx_file_zh = r'e:\pythonProjects\SCI_GO\docs\脑卒中相关性肺炎_SAP_1比4嵌套病例对照研究方案.docx'
    create_styled_document(md_file, docx_file_en)
    create_styled_document(md_file, docx_file_zh)
    print(f"[Success] Generated Word documents successfully.")
